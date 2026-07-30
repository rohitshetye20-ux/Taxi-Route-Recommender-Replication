"""
cover.py

Professional Cover Page
Report Builder Framework v2.0
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches


from report_builder.config import (
    REPORT_TITLE,
    AUTHOR,
    VERSION,
)


def add_cover_page(document):
    """
    Generate the report cover page.
    """

    # -----------------------------------------------------
    # Blank space
    # -----------------------------------------------------

    document.add_paragraph()

    document.add_paragraph()

    document.add_paragraph()

    document.add_paragraph()

    document.add_paragraph()

    # -----------------------------------------------------
    # Title
    # -----------------------------------------------------

    title = document.add_paragraph()

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run(REPORT_TITLE)

    run.bold = True

    run.font.size = Pt(24)

    # -----------------------------------------------------
    # Subtitle
    # -----------------------------------------------------

    document.add_paragraph()

    subtitle = document.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = subtitle.add_run(
        "Technical Reproduction Report"
    )

    run.italic = True

    run.font.size = Pt(16)

    # -----------------------------------------------------
    # Spacer
    # -----------------------------------------------------

    for _ in range(8):

        document.add_paragraph()

    # -----------------------------------------------------
    # Information Table
    # -----------------------------------------------------

    info = document.add_table(rows=4, cols=2)

    info.style = "Table Grid"

    info.cell(0, 0).text = "Author"

    info.cell(0, 1).text = AUTHOR

    info.cell(1, 0).text = "Framework"

    info.cell(1, 1).text = "Final Report Builder"

    info.cell(2, 0).text = "Version"

    info.cell(2, 1).text = VERSION

    info.cell(3, 0).text = "Generated"

    info.cell(3, 1).text = "Automatically"

    # -----------------------------------------------------
    # Page Break
    # -----------------------------------------------------

    document.add_page_break()

    