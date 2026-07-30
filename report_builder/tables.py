"""
tables.py

Markdown Table Renderer
Report Builder Framework v2.0
"""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report_builder.config import TABLE_STYLE


class TableRenderer:
    """
    Converts Markdown tables into Microsoft Word tables.
    """

    def __init__(self, document):

        self.document = document

        self.table_number = 1

    # ======================================================
    # PUBLIC
    # ======================================================

    def render(self, table_lines):
        """
        Render a markdown table into Word.
        """

        if not table_lines:

            return

        # --------------------------------------------------
        # Clean table
        # --------------------------------------------------

        rows = []

        for line in table_lines:

            line = line.strip()

            if not line:
                continue

            rows.append(line)

        if len(rows) < 2:
            return

        # --------------------------------------------------
        # Header
        # --------------------------------------------------

        headers = [
            item.strip()
            for item in rows[0].strip("|").split("|")
        ]

        table = self.document.add_table(
            rows=1,
            cols=len(headers)
        )

        table.style = TABLE_STYLE

        table.autofit = True

        header_cells = table.rows[0].cells

        for i, value in enumerate(headers):

            paragraph = header_cells[i].paragraphs[0]

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run(value)

            run.bold = True

            run.font.size = Pt(11)

        # --------------------------------------------------
        # Body
        # --------------------------------------------------

        for row in rows[2:]:

            values = [
                value.strip()
                for value in row.strip("|").split("|")
            ]

            cells = table.add_row().cells

            for i in range(min(len(values), len(cells))):

                paragraph = cells[i].paragraphs[0]

                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                run = paragraph.add_run(values[i])

                run.font.size = Pt(11)

        # --------------------------------------------------
        # Space after table
        # --------------------------------------------------

        self.document.add_paragraph()

    # ======================================================
    # OPTIONAL TABLE CAPTION
    # ======================================================

    def add_caption(self, caption):

        paragraph = self.document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.style = "Caption"

        run = paragraph.add_run(

            f"Table {self.table_number}. {caption}"

        )

        run.bold = True

        self.table_number += 1