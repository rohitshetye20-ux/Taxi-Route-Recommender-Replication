"""
headers.py

Professional Header and Footer
Report Builder Framework v2.0
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from report_builder.config import (
    REPORT_TITLE,
    VERSION,
)


# ==========================================================
# PAGE NUMBER
# ==========================================================

def add_page_number(paragraph):
    """
    Insert an automatic page number field.
    """

    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "1"

    fld_sep.append(fld_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


# ==========================================================
# HEADER & FOOTER
# ==========================================================

def apply_headers(document):
    """
    Apply professional headers and footers.
    """

    for section in document.sections:

        # --------------------------------------------------
        # HEADER
        # --------------------------------------------------

        header = section.header

        header_para = header.paragraphs[0]

        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header_para.clear()

        run = header_para.add_run(
            f"{REPORT_TITLE}    |    Version {VERSION}"
        )

        run.font.size = Pt(9)

        run.bold = True

        # --------------------------------------------------
        # FOOTER
        # --------------------------------------------------

        footer = section.footer

        footer_para = footer.paragraphs[0]

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_para.clear()

        run = footer_para.add_run("Page ")

        run.font.size = Pt(9)

        add_page_number(footer_para)

        