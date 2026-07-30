"""
builder.py

Main Report Builder
Report Builder Framework v2.0
"""

from docx import Document

from report_builder.config import (
    CHAPTERS,
    CHAPTER_DIR,
    OUTPUT_DIR,
    DOCX_NAME,
)

from report_builder.styles import (
    apply_styles,
    apply_page_layout,
)

from report_builder.headers import (
    apply_headers,
)

from report_builder.cover import (
    add_cover_page,
)

from report_builder.toc import (
    add_table_of_contents,
)

from report_builder.parser import (
    MarkdownParser,
)

from report_builder.exporter import (
    ReportExporter,
)

from report_builder.utils import (
    print_banner,
)


class ReportBuilder:
    """
    Main controller for building the report.
    """

    # ======================================================
    # INITIALIZATION
    # ======================================================

    def __init__(self):

        # Create a new Word document
        self.document = Document()

        # Apply document formatting
        apply_styles(self.document)

        apply_page_layout(self.document)

        apply_headers(self.document)

        # Initialize parser
        self.parser = MarkdownParser(self.document)

    # ======================================================
    # BUILD REPORT
    # ======================================================

    def build(self):

        print_banner("BUILDING FINAL REPORT")

        # ------------------------------------------
        # Cover Page
        # ------------------------------------------

        add_cover_page(self.document)

        # ------------------------------------------
        # Table of Contents
        # ------------------------------------------

        add_table_of_contents(self.document)

        # ------------------------------------------
        # Chapters
        # ------------------------------------------

        for index, chapter in enumerate(CHAPTERS):

            chapter_path = CHAPTER_DIR / chapter

            print(f"Loading : {chapter}")

            self.parser.parse_file(chapter_path)

            # Every chapter except the last ends with a page break
            if index < len(CHAPTERS) - 1:

                self.document.add_page_break()

        # ------------------------------------------
        # Export
        # ------------------------------------------

        exporter = ReportExporter(self.document)

        output_file = OUTPUT_DIR / DOCX_NAME

        return exporter.save_docx(output_file)

    # ======================================================
    # SUMMARY
    # ======================================================

    def summary(self):

        print()

        print("=" * 60)

        print("REPORT BUILDER")

        print("=" * 60)

        print(f"Total Chapters : {len(CHAPTERS)}")

        print(f"Output File    : {DOCX_NAME}")

        print(f"Output Folder  : {OUTPUT_DIR}")

        print("=" * 60)