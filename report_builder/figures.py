"""
figures.py

Automatic Figure Renderer
Report Builder Framework v2.0
"""

from pathlib import Path

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report_builder.config import (
    FIGURE_DIR,
    DEFAULT_IMAGE_WIDTH,
    CAPTION_PREFIX,
    FIGURE_MAP,
)


class FigureRenderer:
    """
    Inserts figures into the Word document.
    """

    def __init__(self, document):

        self.document = document

        self.figure_number = 1

    # ---------------------------------------------------------

    def render(self, placeholder):
        """
        Render a figure into the document.
        """

        # ------------------------------------------
        # Clean placeholder
        # ------------------------------------------

        placeholder = (
            placeholder
            .replace("Diagram", "")
            .replace("Figure", "")
            .strip()
        )

        print(f"\nLooking for figure : {placeholder}")

        image_path = None

        # ------------------------------------------
        # First try FIGURE_MAP
        # ------------------------------------------

        if placeholder in FIGURE_MAP:

            filename = FIGURE_MAP[placeholder]

            candidate = FIGURE_DIR / filename

            print(f"Mapped to : {filename}")

            if candidate.exists():

                image_path = candidate

                print("Status : FOUND")

            else:

                print("Status : FILE NOT FOUND")

        # ------------------------------------------
        # Fallback search
        # ------------------------------------------

        if image_path is None:

            candidates = [

                f"{placeholder}.png",
                f"{placeholder}.jpg",
                f"{placeholder}.jpeg",

                placeholder.lower() + ".png",
                placeholder.lower() + ".jpg",

                placeholder.replace(" ", "_") + ".png",
                placeholder.replace(" ", "_") + ".jpg",

                placeholder.lower().replace(" ", "_") + ".png",
                placeholder.lower().replace(" ", "_") + ".jpg",

            ]

            for filename in candidates:

                candidate = FIGURE_DIR / filename

                if candidate.exists():

                    image_path = candidate

                    print(f"Fallback matched : {filename}")

                    break

        # ------------------------------------------
        # Missing Figure
        # ------------------------------------------

        if image_path is None:

            print(f"Missing Figure : {placeholder}")

            paragraph = self.document.add_paragraph()

            paragraph.style = "Intense Quote"

            paragraph.add_run(
                f"[Missing Figure: {placeholder}]"
            )

            return

        # ------------------------------------------
        # Insert Figure
        # ------------------------------------------

        paragraph = self.document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()

        run.add_picture(
            str(image_path),
            width=Inches(DEFAULT_IMAGE_WIDTH)
        )

        # ------------------------------------------
        # Caption
        # ------------------------------------------

        caption = self.document.add_paragraph()

        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption.style = "Caption"

        caption.add_run(
            f"{CAPTION_PREFIX} {self.figure_number}. {placeholder}"
        ).bold = True

        self.figure_number += 1

        