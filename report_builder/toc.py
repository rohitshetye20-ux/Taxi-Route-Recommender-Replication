"""
toc.py

Automatic Table of Contents
Report Builder Framework v2.0
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ==========================================================
# TABLE OF CONTENTS
# ==========================================================

def add_table_of_contents(document):
    """
    Insert a Microsoft Word automatic Table of Contents.
    """

    # ------------------------------------------------------
    # Title
    # ------------------------------------------------------

    heading = document.add_heading(
        "Table of Contents",
        level=1
    )

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------
    # TOC Field
    # ------------------------------------------------------

    paragraph = document.add_paragraph()

    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")

    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = (
        "Right-click here and choose 'Update Field' "
        "or press Ctrl+A then F9 in Microsoft Word."
    )

    fld_separate.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)

    # ------------------------------------------------------
    # Start Chapter 1 on a new page
    # ------------------------------------------------------

    document.add_page_break()

    